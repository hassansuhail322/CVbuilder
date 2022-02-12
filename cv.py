from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


document = Document()

# document.add_picture('.\desktop\pics\img2.jpg', width=Inches(1))


name=input('enter your name ')
phone=input('enter your phone number :')
email=input('enter your mail id ')
location=  'Delhi, India'#input('enter your state you belongs to ')

document.add_heading(name)
document.add_paragraph("( "+ phone +" )" +'  |  '+ email+ '  |  '+ location)
about=input('tell us about yourself  :')
document.add_paragraph(about)


#                                                 #   word experience 

document.add_heading('PROFRSSIONAL EXPERIENCE ')


n=input("enter number of company you worked in ")
n=int(n)
while n!=0:
    company = input('enter company name ')
    position = input('enter the position in ' + company)
    z=document.add_paragraph()
    z.add_run(position + ' | ' + company).bold=True
    points = input('explain your experience: ')
    document.add_paragraph(points).style = 'List Bullet'

    while True:
        has_more = input('do you want to add more points yes or no  ')
        if has_more == 'yes':
            points = input('explain your experience: ')
            document.add_paragraph(points).style = 'List Bullet'
        else:
            break
    n=n-1
                                            #  skills

document.add_heading("TECHNICAL SKILLS")
language=input('enter the programming language you know : ')
database=input('enter the data base you worked in ')
tech=input('enter the technology you worked in : ')

l=document.add_paragraph()
l.add_run('LANGUAGE: ').bold= True
l.style='List Bullet'
l.add_run(language  )

d=document.add_paragraph()
d.add_run('Database:').bold=True
d.style='List Bullet'
d.add_run(database)

t=document.add_paragraph()
t.add_run('Technologies').bold=True
t.style='List Bullet'
t.add_run(tech)

                                               #  PROJECTS::

document.add_heading('PROJECTS ')
t=input('enter no of projects you want to add in your cv ')
t=int(t)
while t:

    first=input('enter what you have made ')
    second=input('enter what tools and skills you used ')

    p=document.add_paragraph()
    p.add_run(first).bold=True
    p.add_run('  |  ').bold=True
    p.add_run(second)

    explain=input('explain about your project')
    j=document.add_paragraph()
    j.add_run('  '+explain)
    j.style='List Bullet'
    t=t-1


document.save('cv.docx')

# # run = document.add_paragraph() .add_run('some text')
# # font = run.font
# font.color.rgb = RGBColor(255, 0, 0)
# # while projects!=0:
 













